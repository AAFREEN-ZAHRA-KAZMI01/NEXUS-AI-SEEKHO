"""Generate sample.pdf and sample.docx for the test suite."""
import sys
sys.path.insert(0, "..")


def generate_pdf():
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Quarterly Business Report - Q3 2024", ln=True)
        pdf.set_font("Helvetica", size=12)
        pdf.ln(5)
        pdf.multi_cell(
            0, 8,
            "Executive Summary: Orders in Lahore region declined by 25% in Q3 2024, "
            "representing a revenue impact of PKR 2,123,500. Fuel cost increases of "
            "18% are the primary driver."
        )
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Key Metrics:", ln=True)
        pdf.set_font("Helvetica", size=11)
        pdf.cell(0, 8, "Order Volume: 890 units (down from 1,240)", ln=True)
        pdf.cell(0, 8, "Revenue: PKR 6,096,500 (down from PKR 8,500,000)", ln=True)
        pdf.cell(0, 8, "Delivery Cost: PKR 352 per shipment", ln=True)
        pdf.cell(0, 8, "Affected Region: Lahore, Punjab", ln=True)
        pdf.output("tests/fixtures/sample.pdf")
        print("sample.pdf generated")
    except ImportError:
        print("fpdf2 not installed — pip install fpdf2 — skipping PDF generation")


def generate_docx():
    from docx import Document
    doc = Document()
    doc.add_heading("Policy Impact Report: Fuel Price Notification", 0)
    doc.add_heading("Background", 1)
    doc.add_paragraph(
        "OGRA (Oil and Gas Regulatory Authority) issued Notification No. OGRA-2024-1101 "
        "on November 1st, 2024, mandating a PKR 14.97 per litre increase in HSD fuel prices. "
        "This represents an 18.5% increase from the previous rate of PKR 81.02 per litre."
    )
    doc.add_heading("Financial Impact", 1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Before", "After"
    row1 = table.rows[1].cells
    row1[0].text, row1[1].text, row1[2].text = "Fuel Price (PKR/L)", "81.02", "95.99"
    row2 = table.rows[2].cells
    row2[0].text, row2[1].text, row2[2].text = "Delivery Cost (PKR/kg)", "320.00", "378.00"
    row3 = table.rows[3].cells
    row3[0].text, row3[1].text, row3[2].text = "Monthly Cost Impact (PKR)", "1,344,000", "1,587,600"
    doc.add_paragraph("")
    run = doc.add_paragraph().add_run(
        "Recommended Action: Update delivery pricing by +8% immediately."
    )
    run.bold = True
    doc.save("tests/fixtures/sample.docx")
    print("sample.docx generated")


if __name__ == "__main__":
    generate_pdf()
    generate_docx()
