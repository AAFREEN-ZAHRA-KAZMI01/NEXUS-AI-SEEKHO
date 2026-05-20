import io
from docx import Document

from parsers.document_intelligence import extract_document_intelligently


async def parse_docx(file_bytes: bytes, domain: str = "business", use_vision: bool = True) -> dict:
    if use_vision:
        try:
            result = await extract_document_intelligently(file_bytes, "docx", domain, use_vision=True)
            if result and result.get("clean_text"):
                return result
        except Exception as e:
            print(f"[docx_parser] Vision extraction failed, falling back to basic: {e}")

    # Basic fallback with python-docx
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text_parts = []
        headings_list = []
        tables_list = []
        bold_list = []

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                continue

            if para.style.name.startswith("Heading 1"):
                text = f"# {text}"
                headings_list.append(para.text)
            elif para.style.name.startswith("Heading 2"):
                text = f"## {text}"
                headings_list.append(para.text)

            styled_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold:
                    bold_list.append(run_text)
                    styled_text += f"**{run_text}**"
                else:
                    styled_text += run_text

            full_text_parts.append(styled_text)

        for table in doc.tables:
            md_table = []
            for row in table.rows:
                row_str = " | ".join(cell.text.strip() for cell in row.cells)
                md_table.append(f"| {row_str} |")
            if md_table:
                md_str = "\n".join(md_table)
                tables_list.append(md_str)
                full_text_parts.append(f"\n[Table]\n{md_str}\n")

        full_text = "\n\n".join(full_text_parts)
        return {
            "clean_text": full_text[:8000],
            "headings": headings_list,
            "tables_extracted": tables_list,
            "bold_passages": bold_list,
            "source_type": "docx",
            "extraction_method": "basic_ocr",
        }
    except Exception as e:
        return {"error": True, "reason": str(e), "clean_text": ""}
